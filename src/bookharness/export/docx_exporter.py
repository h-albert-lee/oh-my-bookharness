"""Export Markdown drafts to Word (.docx) using the publisher template.

Uses the 한빛미디어 template (1장_피드백 진행.docx) as a style source,
mapping Markdown elements to the publisher's custom Word styles.

Style mapping:
    # N장. 제목          → Heading 1
    ## N.M 절제목         → Heading 2
    ### N.M.K 중제목      → Heading 3
    #### 소제목           → Heading 4
    ##### 소소제목        → 소소제목 (custom)
    ::: box 제목 ... :::  → 박스 제목 + 박스_끝
    ::: note 제목 ... ::: → 노트
    ::: tip 제목 ... :::  → 노트 (same style as note)
    ```code```            → _코드 (custom paragraph style)
    `inline code`         → _코드_본문 (custom character style)
    **표 N-M. 제목**      → Caption
    **그림 N-M. 제목**    → Caption
    - list items          → List Paragraph
    [^source_id]          → Word footnote
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


class DocxExporter:
    """Convert a Markdown draft to .docx using the publisher template."""

    def __init__(self, template_path: Path | str | None = None,
                 source_bundle: dict | None = None) -> None:
        self.template_path = Path(template_path) if template_path else None
        self._source_bundle = source_bundle or {}
        self._footnote_id = 1  # auto-increment footnote IDs
        self._footnotes_xml = None  # lxml element for footnotes part
        self._footnotes_part = None  # the OPC part object
        self._created_footnotes: dict[str, str] = {}  # ref_id -> fn_id (dedup)

    def export(self, markdown_path: Path | str, output_path: Path | str | None = None) -> Path:
        markdown_path = Path(markdown_path)
        if output_path is None:
            output_path = markdown_path.with_suffix(".docx")
        else:
            output_path = Path(output_path)

        md_text = markdown_path.read_text(encoding="utf-8")

        if self.template_path and self.template_path.exists():
            doc = Document(str(self.template_path))
            self._clear_content(doc)
        else:
            doc = Document()

        self._footnote_id = 1
        self._created_footnotes = {}
        self._init_footnotes(doc)
        lines = md_text.split("\n")
        self._render(doc, lines)

        # Write back modified footnotes XML
        self._save_footnotes()

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _clear_content(self, doc: Document) -> None:
        for _ in range(len(doc.paragraphs)):
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        for _ in range(len(doc.tables)):
            t = doc.tables[0]._element
            t.getparent().remove(t)

    def _build_source_citation(self, source_id: str) -> str | None:
        """Build a citation string from the source bundle metadata."""
        for key in ("core_sources", "supporting_sources"):
            for s in self._source_bundle.get(key, []):
                sid = s.get("source_id", "")
                if sid == source_id:
                    title = s.get("title", "")
                    url = s.get("url", "")
                    authors = s.get("authors", "")
                    year = s.get("year", "")
                    parts = []
                    if authors:
                        parts.append(str(authors))
                    if year:
                        parts.append(f"({year})")
                    if title:
                        parts.append(title)
                    if url:
                        parts.append(url)
                    return ". ".join(parts) if parts else None
        return None

    def _render(self, doc: Document, lines: list[str]) -> None:
        i = 0
        footnotes: dict[str, str] = {}

        # First pass: collect footnote definitions [^id]: text
        for line in lines:
            m = re.match(r"^\[\^(\w+)\]:\s*(.+)$", line)
            if m:
                footnotes[m.group(1)] = m.group(2)

        # Second pass: fill missing definitions from source bundle
        all_refs = set(re.findall(r"\[\^(\w+)\](?!:)", "\n".join(lines)))
        for ref_id in all_refs:
            if ref_id not in footnotes:
                citation = self._build_source_citation(ref_id)
                if citation:
                    footnotes[ref_id] = citation

        while i < len(lines):
            line = lines[i]

            if not line.strip():
                i += 1
                continue

            # Skip footnote definitions
            if re.match(r"^\[\^(\w+)\]:\s*", line):
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^---+\s*$", line):
                i += 1
                continue

            # Fenced div: ::: box/note/tip
            m = re.match(r"^:::\s*(box|note|tip)\s*(.*)?$", line)
            if m:
                i = self._render_fenced_div(doc, lines, i, m.group(1), m.group(2) or "", footnotes)
                continue

            # Code block
            if line.startswith("```"):
                i = self._render_code_block(doc, lines, i)
                continue

            # Headings
            if line.startswith("#"):
                self._render_heading(doc, line)
                i += 1
                continue

            # Table caption (**표 N-M.** or **그림 N-M.**)
            if re.match(r"^\*\*(?:표|그림)\s+\d+-\d+\.\s*.+\*\*$", line.strip()):
                caption_text = line.strip().strip("*")
                self._add_para(doc, caption_text, "Caption")
                i += 1
                continue

            # Markdown table
            if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1].strip()):
                i = self._render_table(doc, lines, i)
                continue

            # List items
            if re.match(r"^[-*]\s+", line):
                text = re.sub(r"^[-*]\s+", "", line)
                para = self._add_para(doc, None, "List Paragraph")
                self._add_formatted_runs(para, text, footnotes, doc)
                i += 1
                continue

            # Numbered list
            m = re.match(r"^(\d+)[.)]\s+", line)
            if m:
                text = line[m.end():]
                para = self._add_para(doc, None, "List Paragraph")
                self._add_formatted_runs(para, text, footnotes, doc)
                i += 1
                continue

            # Normal paragraph
            para = self._add_para(doc, None, "Normal")
            self._add_formatted_runs(para, line, footnotes, doc)
            i += 1

    def _render_heading(self, doc: Document, line: str) -> None:
        level = 0
        for ch in line:
            if ch == "#":
                level += 1
            else:
                break
        text = line[level:].strip()

        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
        style = style_map.get(level, "소소제목")
        self._add_para(doc, text, style)

    def _render_fenced_div(self, doc: Document, lines: list[str], start: int,
                           div_type: str, title: str, footnotes: dict) -> int:
        i = start + 1
        content_lines = []

        while i < len(lines):
            if lines[i].strip() == ":::":
                i += 1
                break
            content_lines.append(lines[i])
            i += 1

        title = title.strip()

        if div_type == "box":
            if title:
                self._add_para(doc, title, "박스 제목")
            for para_text in self._split_paragraphs(content_lines):
                para = self._add_para(doc, None, "박스_끝")
                self._add_formatted_runs(para, para_text, footnotes, doc)
        else:
            # Note/Tip
            if title:
                self._add_para(doc, title, "노트")
            for para_text in self._split_paragraphs(content_lines):
                para = self._add_para(doc, None, "노트")
                self._add_formatted_runs(para, para_text, footnotes, doc)

        return i

    def _render_code_block(self, doc: Document, lines: list[str], start: int) -> int:
        i = start + 1
        code_lines = []

        while i < len(lines):
            if lines[i].startswith("```"):
                i += 1
                break
            code_lines.append(lines[i])
            i += 1

        for code_line in code_lines:
            self._add_para(doc, code_line, "_코드")

        return i

    def _render_table(self, doc: Document, lines: list[str], start: int) -> int:
        i = start
        rows_data = []

        while i < len(lines) and "|" in lines[i]:
            row_text = lines[i].strip()
            if re.match(r"^\|[-\s|:]+\|$", row_text):
                i += 1
                continue
            cells = [c.strip() for c in row_text.split("|")[1:-1]]
            rows_data.append(cells)
            i += 1

        if not rows_data:
            return i

        num_cols = max(len(row) for row in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = "Table Grid"

        for ri, row in enumerate(rows_data):
            for ci, cell_text in enumerate(row):
                if ci < num_cols:
                    cell_text = self._strip_md_formatting(cell_text)
                    table.rows[ri].cells[ci].text = cell_text

        return i

    def _split_paragraphs(self, lines: list[str]) -> list[str]:
        paragraphs = []
        current = []
        for line in lines:
            if not line.strip():
                if current:
                    paragraphs.append(" ".join(current))
                    current = []
            else:
                current.append(line.strip())
        if current:
            paragraphs.append(" ".join(current))
        return paragraphs

    def _strip_md_formatting(self, text: str) -> str:
        """Remove all markdown inline formatting markers from text."""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
        text = re.sub(r"\*(.+?)\*", r"\1", text)       # *italic*
        text = re.sub(r"`([^`]+)`", r"\1", text)       # `code`
        text = re.sub(r"\[\^(\w+)\]", "", text)         # [^footnote]
        return text

    def _add_para(self, doc: Document, text: str | None, style_name: str):
        """Add a paragraph with the given style.

        If text is None, returns an empty paragraph for caller to fill with runs.
        If text is a string, adds it as a plain run (for headings, captions, code).
        """
        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph()
        if text is not None:
            para.add_run(text)
        return para

    def _add_formatted_runs(self, para, text: str, footnotes: dict[str, str], doc: Document) -> None:
        """Parse inline Markdown formatting and add Word runs with proper styles."""
        # Match: **bold**, `code`, [^footnote]
        pattern = re.compile(
            r"(\*\*(.+?)\*\*)"       # bold
            r"|(\*(.+?)\*)"           # italic
            r"|(`([^`]+)`)"           # inline code
            r"|(\[\^(\w+)\])"         # footnote reference
        )

        pos = 0
        for m in pattern.finditer(text):
            # Add plain text before this match
            if m.start() > pos:
                para.add_run(text[pos:m.start()])

            if m.group(2):  # **bold**
                run = para.add_run(m.group(2))
                run.bold = True
            elif m.group(4):  # *italic*
                run = para.add_run(m.group(4))
                run.italic = True
            elif m.group(6):  # `inline code`
                run = para.add_run(m.group(6))
                try:
                    run.style = doc.styles["_코드_본문"]
                except KeyError:
                    run.font.name = "Consolas"
            elif m.group(8):  # [^footnote_id]
                ref_id = m.group(8)
                fn_text = footnotes.get(ref_id, "")
                if not fn_text:
                    # Skip refs with no definition at all
                    continue
                if ref_id in self._created_footnotes:
                    # Reuse existing footnote number
                    self._add_footnote_ref(para, self._created_footnotes[ref_id])
                else:
                    fn_id = self._add_footnote(para, fn_text, doc)
                    if fn_id:
                        self._created_footnotes[ref_id] = fn_id

            pos = m.end()

        # Add remaining text after last match
        if pos < len(text):
            para.add_run(text[pos:])

    def _init_footnotes(self, doc: Document) -> None:
        """Load the footnotes XML part from the document."""
        for rel in doc.part.rels.values():
            if "footnote" in str(rel.reltype).lower():
                self._footnotes_part = rel.target_part
                self._footnotes_xml = etree.fromstring(self._footnotes_part.blob)
                # Find max existing footnote ID to avoid collisions
                for fn in self._footnotes_xml:
                    fn_id = fn.get(qn("w:id"))
                    if fn_id and fn_id.lstrip("-").isdigit():
                        self._footnote_id = max(self._footnote_id, int(fn_id) + 1)
                return
        self._footnotes_xml = None
        self._footnotes_part = None

    def _save_footnotes(self) -> None:
        """Write modified footnotes XML back to the part."""
        if self._footnotes_part is not None and self._footnotes_xml is not None:
            self._footnotes_part._blob = etree.tostring(self._footnotes_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    def _add_footnote_ref(self, paragraph, fn_id: str) -> None:
        """Add a reference to an existing footnote (for duplicate refs)."""
        ref_run = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        ref_rstyle = OxmlElement("w:rStyle")
        ref_rstyle.set(qn("w:val"), "footnote reference")
        ref_rpr.append(ref_rstyle)
        ref_run.append(ref_rpr)
        fn_ref = OxmlElement("w:footnoteReference")
        fn_ref.set(qn("w:id"), fn_id)
        ref_run.append(fn_ref)
        paragraph._element.append(ref_run)

    def _add_footnote(self, paragraph, footnote_text: str, doc: Document) -> str | None:
        """Add a real Word footnote to the paragraph. Returns fn_id."""
        if self._footnotes_xml is None:
            run = paragraph.add_run("*")
            run.font.superscript = True
            return None

        fn_id = str(self._footnote_id)
        self._footnote_id += 1

        # Build footnote element
        footnote_el = etree.SubElement(self._footnotes_xml, qn("w:footnote"))
        footnote_el.set(qn("w:id"), fn_id)

        fn_para = etree.SubElement(footnote_el, qn("w:p"))

        # Paragraph style: footnote text
        fn_ppr = etree.SubElement(fn_para, qn("w:pPr"))
        fn_pstyle = etree.SubElement(fn_ppr, qn("w:pStyle"))
        fn_pstyle.set(qn("w:val"), "footnote text")

        # Footnote ref mark (auto-number inside the footnote)
        fn_ref_run = etree.SubElement(fn_para, qn("w:r"))
        fn_ref_rpr = etree.SubElement(fn_ref_run, qn("w:rPr"))
        fn_ref_rstyle = etree.SubElement(fn_ref_rpr, qn("w:rStyle"))
        fn_ref_rstyle.set(qn("w:val"), "footnote reference")
        etree.SubElement(fn_ref_run, qn("w:footnoteRef"))

        # Space
        space_run = etree.SubElement(fn_para, qn("w:r"))
        space_t = etree.SubElement(space_run, qn("w:t"))
        space_t.set(qn("xml:space"), "preserve")
        space_t.text = " "

        # Footnote body text
        text_run = etree.SubElement(fn_para, qn("w:r"))
        text_t = etree.SubElement(text_run, qn("w:t"))
        text_t.set(qn("xml:space"), "preserve")
        text_t.text = footnote_text

        # Insert reference mark in the main document paragraph
        ref_run = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        ref_rstyle = OxmlElement("w:rStyle")
        ref_rstyle.set(qn("w:val"), "footnote reference")
        ref_rpr.append(ref_rstyle)
        ref_run.append(ref_rpr)
        fn_ref = OxmlElement("w:footnoteReference")
        fn_ref.set(qn("w:id"), fn_id)
        ref_run.append(fn_ref)
        paragraph._element.append(ref_run)
        return fn_id


def _load_bundle(project_root: Path, chapter_id: str) -> dict:
    """Load the source bundle for a chapter."""
    bundle_path = project_root / f"sources/chapter_packs/{chapter_id}/bundle.yaml"
    if bundle_path.exists():
        import yaml
        return yaml.safe_load(bundle_path.read_text(encoding="utf-8")) or {}
    return {}


def export_chapter(
    project_root: Path,
    chapter_id: str,
    template_path: Path | str | None = None,
    output_dir: Path | str | None = None,
) -> Path:
    chapter_dir = project_root / f"manuscript/{chapter_id}"
    drafts = sorted(chapter_dir.glob("draft_v*.md"))
    if not drafts:
        raise FileNotFoundError(f"No drafts found for {chapter_id}")

    latest_draft = drafts[-1]

    if output_dir is None:
        output_dir = project_root / "export"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / f"{chapter_id}_{latest_draft.stem}.docx"

    bundle = _load_bundle(project_root, chapter_id)
    exporter = DocxExporter(template_path, source_bundle=bundle)
    return exporter.export(latest_draft, output_path)
